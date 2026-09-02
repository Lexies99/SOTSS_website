import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, text, title="NOTE:", fill_color="EEF6FF", border_color="0B2F5B"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, fill_color)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Left border only
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"{title} ")
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(11, 47, 91)
    r_title.font.size = Pt(10)
    
    r_text = p.add_run(text)
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph()

def build_docx():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Header & Footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("GIMPA SOTSS — Department Web & Intranet User Manual  |  Page ")
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(100, 116, 139)

    # Base Colors
    NAVY = RGBColor(11, 47, 91)      # #0B2F5B
    BLUE_ACCENT = RGBColor(2, 132, 199) # #0284C7
    SLATE = RGBColor(71, 85, 105)    # #475569
    DARK_TEXT = RGBColor(15, 23, 42)  # #0F172A

    # Document Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("GHANA INSTITUTE OF MANAGEMENT AND PUBLIC ADMINISTRATION\n")
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = SLATE

    r2 = title_p.add_run("SCHOOL OF TECHNOLOGY & SOCIAL SCIENCES (SOTSS)\n")
    r2.font.size = Pt(13)
    r2.bold = True
    r2.font.color.rgb = BLUE_ACCENT

    r3 = title_p.add_run("DEPARTMENT OF COMPUTER SCIENCE & INFORMATION SYSTEMS")
    r3.font.size = Pt(16)
    r3.bold = True
    r3.font.color.rgb = NAVY

    doc.add_paragraph()

    # Banner Title
    h_main = doc.add_paragraph()
    h_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_main.paragraph_format.space_before = Pt(12)
    h_main.paragraph_format.space_after = Pt(18)
    r_main = h_main.add_run("System User Manual & Portal Operating Guide")
    r_main.font.size = Pt(22)
    r_main.bold = True
    r_main.font.color.rgb = NAVY

    # Metadata Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    r_sub = sub_p.add_run("A Complete Guide for Prospective Students, Current Students, Faculty Members, Researchers & System Administrators\nDocument Release: August 2026 | Version 2.0")
    r_sub.font.size = Pt(9.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = SLATE

    doc.add_heading("Table of Contents", level=1)
    
    toc_items = [
        ("1. Executive Summary & System Architecture", "3"),
        ("2. Public Portal Features & User Guide", "4"),
        ("    2.1 Header Navigation & Real-Time Site Search", "4"),
        ("    2.2 Home Page Overview", "4"),
        ("    2.3 Department Introduction & History", "5"),
        ("    2.4 Academic Programmes Directory (12 Degree Specs)", "5"),
        ("    2.5 Research Strengths & Focus Areas", "6"),
        ("    2.6 Contact Directory & Interactive Faculty Filter", "7"),
        ("    2.7 News Feed & Social Media Channels", "7"),
        ("    2.8 Cooperation, Industry Partnerships & Alumni", "8"),
        ("3. Faculty & Lecturer Intranet Operating Guide", "8"),
        ("    3.1 Account Registration & GIMPA Email Domain Verification", "8"),
        ("    3.2 Secure Login & Session Security", "9"),
        ("    3.3 Profile Overview & Research Project Management", "9"),
        ("    3.4 Publication & Verification Center (Crawled Alerts vs. Manual)", "10"),
        ("    3.5 Messages Box & School Announcements", "11"),
        ("    3.6 Account Security & Password Modification", "11"),
        ("4. System Administrator Operating Guide", "12"),
        ("    4.1 News Moderation & Crawler Management", "12"),
        ("    4.2 Faculty Profile & Department-Wide Research Audit", "12"),
        ("5. Troubleshooting & Support", "13")
    ]
    
    for item, pg in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25 if item.startswith("    ") else 0)
        r1 = p.add_run(item)
        r1.font.size = Pt(10)
        r1.font.color.rgb = DARK_TEXT
        if not item.startswith("    "):
            r1.bold = True

    doc.add_page_break()

    # SECTION 1
    h1 = doc.add_heading("1. Executive Summary & System Architecture", level=1)
    h1.runs[0].font.color.rgb = NAVY

    p = doc.add_paragraph(
        "The Department of Computer Science & Information Systems Web Portal and Intranet System is an integrated, "
        "high-performance Single Page Application (SPA) designed to serve as the digital hub for academic programmes, "
        "cutting-edge research, faculty publications, student admissions, and internal department communication at the "
        "Ghana Institute of Management and Public Administration (GIMPA)."
    )
    p.runs[0].font.size = Pt(10.5)

    p2 = doc.add_paragraph(
        "Formed through the strategic consolidation of the Department of Computer Sciences and the Department of Information Systems "
        "and Innovation under the School of Technology and Social Sciences (SOTSS), the department provides rigorous undergraduate, "
        "postgraduate diploma, master's, and doctoral education."
    )
    p2.runs[0].font.size = Pt(10.5)

    doc.add_heading("Target User Roles", level=2)
    roles_table = doc.add_table(rows=5, cols=2)
    roles_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = roles_table.rows[0].cells
    hdr_cells[0].text = "User Role"
    hdr_cells[1].text = "Primary System Access & Capabilities"
    for cell in hdr_cells:
        set_cell_background(cell, "0B2F5B")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    row_data = [
        ("Prospective Students", "Browse 12 academic programmes, review entry requirements (WASSCE, HND, First Degree), explore faculty specializations, and initiate online application via GIMPA Admissions Portal."),
        ("Current Students", "Access course info, LMS links, department news, faculty contact details, student advisory, and industry placement/internship opportunities."),
        ("Faculty & Lecturers", "Access the secure Intranet Portal (#intranet), manage academic profile details, submit/confirm research publications, manage active projects, send direct messages, and broadcast announcements."),
        ("Department Administrators / HOD", "Moderate news articles, review department-wide publication verification metrics, assign administrative roles, manage faculty records, and issue global announcements.")
    ]

    for i, (r_role, r_cap) in enumerate(row_data):
        row_cells = roles_table.rows[i+1].cells
        row_cells[0].text = r_role
        row_cells[1].text = r_cap
        bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        set_cell_background(row_cells[0], bg)
        set_cell_background(row_cells[1], bg)
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    # SECTION 2
    h2 = doc.add_heading("2. Public Portal Features & User Guide", level=1)
    h2.runs[0].font.color.rgb = NAVY

    doc.add_heading("2.1 Header Navigation & Real-Time Site Search", level=2)
    doc.add_paragraph(
        "The top navigation bar provides instant access to essential GIMPA portals and system navigation:\n"
        "• GIMPA Home: Redirects to the main GIMPA institutional portal (gimpa.edu.gh).\n"
        "• LMS: Direct access to the GIMPA Learning Management System (lms.gimpa.edu.gh).\n"
        "• Apply Now: Direct portal link for prospective student applications (apply.gimpa.edu.gh/start).\n"
        "• Library: Link to the GIMPA Digital Library & Research Repositories.\n"
        "• Intranet Portal (#intranet): Displays once authenticated, providing access to lecturer tools.\n"
        "• Notification Bell: Displays unverified publication alerts count for logged-in faculty.\n"
        "• Real-Time Search Bar (#siteSearch): Positioned in the main header, users can type keywords (e.g. 'cybersecurity', 'Engmann', 'BSc') to view dynamic autocomplete matches for pages, faculty profiles, and news articles."
    )

    doc.add_heading("2.2 Home Page Overview (#home)", level=2)
    doc.add_paragraph(
        "The Home section features a dynamic slideshow highlighting campus facilities and department life, "
        "followed by quick navigation cards for Department Introduction, Research & Expertise, Staff & Contact, "
        "and Cooperation. It also showcases the latest department news and a quick-view grid of flagship degree programmes."
    )

    doc.add_heading("2.3 Department Introduction & History (#introduction)", level=2)
    doc.add_paragraph(
        "Learn about GIMPA's history (founded in 1961), the establishment of the Department of Computer Sciences in 2017, "
        "and the landmark merger in September 2023 that created the Department of Computer Science & Information Systems. "
        "This section features the official Welcome Message from the Head of Department (Dr. Felicia Engmann) and outlines the core values: "
        "Excellence, Quality, and Connectedness."
    )

    doc.add_heading("2.4 Academic Programmes Directory", level=2)
    doc.add_paragraph(
        "The department offers 12 comprehensive degree and diploma programmes across undergraduate, postgraduate diploma, master's, and doctoral levels. "
        "Clicking any programme card opens its dedicated detail page (#programme/:id) containing full overview, core curriculum modules, admission criteria, and career prospects:"
    )

    prog_table = doc.add_table(rows=13, cols=4)
    prog_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_hdrs = prog_table.rows[0].cells
    p_hdrs[0].text = "Programme Title"
    p_hdrs[1].text = "Level"
    p_hdrs[2].text = "Duration & Mode"
    p_hdrs[3].text = "Key Focus & Target Career"
    for cell in p_hdrs:
        set_cell_background(cell, "0B2F5B")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)

    programmes_list = [
        ("BSc Computer Science", "Undergraduate", "4 Years (Full-Time)", "Algorithms, AI, Robotics, Software Engineering, Networks"),
        ("BSc Information & Comm. Tech (ICT)", "Undergraduate", "4 Years (Day/Evening)", "Web/Mobile App Dev, IT Governance, Enterprise Systems"),
        ("BSc Health Informatics", "Undergraduate", "4 Years (Full-Time)", "Healthcare Data, EMRs, Telemedicine, Health Privacy"),
        ("PG Diploma in ICT", "Postgrad Diploma", "1 Year (Evening/Weekend)", "Conversion pathway into computing for non-IT graduates"),
        ("PG Diploma in MIS", "Postgrad Diploma", "1 Year (Evening/Weekend)", "Enterprise Data, MIS Foundations, IT Infrastructure"),
        ("MSc/MPhil ICT", "Postgraduate", "2 Years (Weekends/Evening)", "Cloud Computing, Decision Support Systems, Advanced Networks"),
        ("MSc IT & Law", "Postgraduate", "2 Years (Weekends/Evening)", "Cybercrime Law, Digital Forensics Law, Data Protection/GDPR"),
        ("MPhil Management Info Systems", "Postgrad Research", "2 Years (Regular Day)", "Research-intensive degree, IS Theories, PhD preparation"),
        ("MSc Management Info Systems", "Postgrad Executive", "2 Years (Evening/Weekend)", "IT Leadership, ERP Systems, Strategic IT Alignment"),
        ("MSc Digital Forensics & Cybersecurity", "Postgraduate", "1 Year (Weekends)", "Penetration Testing, Reverse Eng, Incident Response, ISO 27001"),
        ("MSc Industrial Analytics", "Postgraduate", "2 Years (Flexible)", "Operations Research, Optimization, Supply Chain Analytics"),
        ("PhD Information Systems", "Doctoral", "3 Years (Full-Time)", "Empirical IS research, Philosophy of Technology, Academia/Policy")
    ]

    for idx, (p_title, p_level, p_dur, p_focus) in enumerate(programmes_list):
        row_c = prog_table.rows[idx+1].cells
        row_c[0].text = p_title
        row_c[1].text = p_level
        row_c[2].text = p_dur
        row_c[3].text = p_focus
        bg = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for c in row_c:
            set_cell_background(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    doc.add_heading("2.5 Research Strengths & Focus Areas (#research)", level=2)
    doc.add_paragraph(
        "The department actively conducts research in Predictive Analytics, Wireless Sensor Networks & IoT, Machine Learning, "
        "Computational Immunology, Cybersecurity & Digital Forensics, and Scientific Computing. "
        "Additionally, the department utilizes Simio simulation software under an academic grant from Simio LLC to model industrial processes."
    )

    doc.add_heading("2.6 Contact Directory & Interactive Faculty Filter (#contact)", level=2)
    doc.add_paragraph(
        "The Contact section includes department location (Greenhill, Accra), phone contacts, and email options. "
        "It features an Interactive Faculty Search Bar (#facultySearchInput) allowing real-time filtering of staff cards by name, "
        "office, role, or specialization. Clicking any faculty card opens their full Faculty Profile (#profile/:id), detailing "
        "their academic background, active research projects, and verified publications."
    )

    doc.add_heading("2.7 Department News Feed (#news)", level=2)
    doc.add_paragraph(
        "The News Feed features official department news, international partnership updates (e.g. UESTC China, York University), "
        "and faculty publications. Features include:\n"
        "• News Relevance Filter: Automatically isolates SOTSS computing and technology news stories.\n"
        "• Pagination Controls: 6 stories per page with smooth navigation.\n"
        "• Social Media Toolbar: Direct links to official SOTSS channels on Instagram (@sotss_gimpa), Facebook, and LinkedIn."
    )

    doc.add_heading("2.8 Cooperation, Industry Partnerships & Alumni (#cooperation)", level=2)
    doc.add_paragraph(
        "Details industry collaboration, student internship placements, guest lectures, international academic exchange, "
        "and alumni engagement opportunities."
    )

    add_callout(doc, "All public pages are accessible without login. Any user can view degree requirements, faculty profiles, and research highlights.", title="PUBLIC ACCESS TIP:")

    doc.add_page_break()

    # SECTION 3
    h3 = doc.add_heading("3. Faculty & Lecturer Intranet Operating Guide", level=1)
    h3.runs[0].font.color.rgb = NAVY

    doc.add_heading("3.1 Account Registration & GIMPA Email Domain Verification (#signup)", level=2)
    doc.add_paragraph(
        "Faculty members and lecturers can register for an Intranet account via the Signup portal (#signup). "
        "To ensure institutional security, strict frontend and backend email domain validation is enforced:"
    )

    add_callout(
        doc,
        "Registration is restricted to official GIMPA email addresses ending in @gimpa.edu.gh, @adj.gimpa.edu.gh, or @st.gimpa.edu.gh. "
        "Registrations using generic webmail providers (Gmail, Yahoo, Hotmail) will be automatically rejected.",
        title="EMAIL DOMAIN REQUIREMENT:",
        fill_color="FEF2F2",
        border_color="DC2626"
    )

    doc.add_paragraph(
        "Registration Steps:\n"
        "1. Navigate to #signup or click 'Sign Up' on the login screen.\n"
        "2. Complete the required fields: Full Name (with Title), GIMPA Email, Department, Academic Role, Office Location, and Specialization.\n"
        "3. Optional fields: Phone Number, Extra Role Label/Value (e.g. 'Coordinator, PhD IS').\n"
        "4. Click 'Submit Registration'. The system generates a secure temporary password sent to your email (and logged in the server console for immediate admin assistance)."
    )

    doc.add_heading("3.2 Secure Login & Session Security (#login)", level=2)
    doc.add_paragraph(
        "Navigate to #login, enter your GIMPA email address and password. Upon verification, the backend issues a JSON Web Token (JWT) "
        "stored securely in browser local storage. The top navigation updates to show 'Welcome, [Name]' and unlocks the Intranet Portal (#intranet)."
    )

    doc.add_heading("3.3 Profile Overview & Research Project Management", level=2)
    doc.add_paragraph(
        "Within the Intranet Dashboard, the Profile Overview tab allows lecturers to:\n"
        "• View Official Record: Check name, email, department, role, office location, and specializations.\n"
        "• Manage Research Projects: View current active projects. Add new projects using the '➕ Add New Project' form (Title, Description, optional Project URL).\n"
        "• Update Profile Details: Edit office location, phone number, specialization list, or extra role descriptors."
    )

    doc.add_heading("3.4 Publication & Verification Center", level=2)
    doc.add_paragraph(
        "The Verification Center forms the core of the department's research audit system, combining automated web crawlers with lecturer verification:"
    )

    doc.add_paragraph(
        "A. Crawled Research Alerts:\n"
        "Background scrapers automatically scan academic databases (Google Scholar, Europe PMC, bioRxiv, PubMed) for publications authored by GIMPA faculty. "
        "Discovered papers appear in the lecturer's 'Crawled Alerts' tab in unverified status. Lecturers review the paper title, journal, year, and abstract, then choose:\n"
        "  - 'Confirm Work': Verifies the publication and publishes it to your public Faculty Profile (#profile/:id).\n"
        "  - 'Not Mine': Rejects the publication, removing it from your alerts box.\n\n"
        "B. Manual Publication Submission:\n"
        "If the automated scraper misses a publication, lecturers can manually submit it using the '➕ Submit Publication Manually' form. "
        "Required inputs include Publication Title, Year, Journal/Conference, Authors, Summary/Abstract, and optional DOI/URL. Manually submitted papers are verified immediately.\n\n"
        "C. Managing Verified Publications:\n"
        "Displays all confirmed papers with full citation details. Lecturers can click '✕ Remove' to unpublish or delete an entry if necessary."
    )

    doc.add_heading("3.5 Messages Box & School Announcements", level=2)
    doc.add_paragraph(
        "The Intranet Messages Box facilitates communication across the department:\n"
        "• Global Announcements: Read school-wide announcements broadcasted by the Head of Department or system administrators.\n"
        "• Direct Messages: Receive and view private messages sent specifically to your user account.\n"
        "• Compose & Send: Lecturers can compose messages and select the scope — either broadcasting a 'School-wide Announcement' or addressing a 'Direct Message' to a specific colleague."
    )

    doc.add_heading("3.6 Account Security & Password Modification", level=2)
    doc.add_paragraph(
        "Lecturers can update their portal password at any time via the 'Change Portal Password' card inside the Intranet. "
        "Enter your Old Password, New Password, and Confirm New Password to apply changes."
    )

    doc.add_page_break()

    # SECTION 4
    h4 = doc.add_heading("4. System Administrator Operating Guide", level=1)
    h4.runs[0].font.color.rgb = NAVY

    doc.add_paragraph(
        "Users logged in with Administrator privileges (such as the Head of Department or IT Administrators) gain access to specialized moderation tools in the Intranet:"
    )

    doc.add_heading("4.1 News Moderation & Crawler Management", level=2)
    doc.add_paragraph(
        "• Trigger Automated Crawler: Run automated scrapers to fetch news from RSS feeds and social media.\n"
        "• Review Draft Articles: Edit titles, descriptions, images, or external links before publishing.\n"
        "• Publish / Unpublish: Control which news stories appear on the public Home page (#home) and News Feed (#news)."
    )

    doc.add_heading("4.2 Faculty Profile & Department-Wide Research Audit", level=2)
    doc.add_paragraph(
        "• Faculty Directory Management: Inspect all registered faculty accounts, assign administrator flags (`is_admin`), and update department assignments.\n"
        "• Department Publication Audit: Inspect total verified vs. unverified publication statistics across the entire department for accreditation and annual reporting."
    )

    # SECTION 5
    h5 = doc.add_heading("5. Troubleshooting & Support", level=1)
    h5.runs[0].font.color.rgb = NAVY

    faq_table = doc.add_table(rows=5, cols=2)
    faq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    f_hdrs = faq_table.rows[0].cells
    f_hdrs[0].text = "Common Issue / Question"
    f_hdrs[1].text = "Recommended Solution & Troubleshooting Workflow"
    for cell in f_hdrs:
        set_cell_background(cell, "0B2F5B")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)

    faqs = [
        ("Unable to register with personal email (Gmail/Yahoo)", "System safety requires a valid GIMPA email address (@gimpa.edu.gh, @adj.gimpa.edu.gh, or @st.gimpa.edu.gh). Use your official institutional email."),
        ("Forgot Intranet Portal Password", "Contact the Department Administrator or HOD at csshead@gimpa.edu.gh to reset your password or retrieve your initial temporary credentials."),
        ("Crawled publication belongs to another author with a similar name", "Click 'Not Mine' in the Verification Center. The paper will be discarded from your profile without affecting other faculty."),
        ("Newly added project or publication not showing on public profile", "Ensure you clicked 'Confirm Work' or 'Add Project' successfully. Refresh your browser or navigate to #profile/[your-id] to verify.")
    ]

    for idx, (q, a) in enumerate(faqs):
        row_c = faq_table.rows[idx+1].cells
        row_c[0].text = q
        row_c[1].text = a
        bg = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for c in row_c:
            set_cell_background(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    # Final Contact Footer Box
    add_callout(
        doc,
        "Department of Computer Science & Information Systems\n"
        "School of Technology & Social Sciences (SOTSS), GIMPA\n"
        "Greenhill, Accra, Ghana  |  Email: csshead@gimpa.edu.gh  |  Phone: +233 (0) 501620138\n"
        "Website: https://gimpa.edu.gh/",
        title="NEED FURTHER ASSISTANCE?",
        fill_color="F1F5F9",
        border_color="0B2F5B"
    )

    out_dir = r"d:\NSS\SOTSS_website\SOTSS_website"
    out_path = os.path.join(out_dir, "SOTSS_Website_User_Manual.docx")
    try:
        doc.save(out_path)
        print(f"Successfully saved Word Document user manual to: {out_path}")
    except PermissionError:
        out_path = os.path.join(out_dir, "SOTSS_Website_User_Manual_v2.docx")
        doc.save(out_path)
        print(f"Warning: Main file locked. Saved Word Document user manual to: {out_path}")

    # Also save to docs subfolder if it exists
    docs_dir = os.path.join(out_dir, "docs")
    if not os.path.exists(docs_dir):
        os.makedirs(docs_dir, exist_ok=True)
    try:
        doc.save(os.path.join(docs_dir, "SOTSS_Website_User_Manual.docx"))
        print(f"Successfully saved copy to docs folder.")
    except PermissionError:
        doc.save(os.path.join(docs_dir, "SOTSS_Website_User_Manual_v2.docx"))
        print(f"Warning: Docs copy locked. Saved to docs/SOTSS_Website_User_Manual_v2.docx")

if __name__ == "__main__":
    build_docx()

